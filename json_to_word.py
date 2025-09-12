import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Load parsed JSON file
with open("ParsedResumeJSON.txt", "r", encoding="utf-8") as f:
    resume_data = json.load(f)

# Load suggestions file if it exists
def load_suggestions():
    try:
        with open("suggestions_output.txt", "r", encoding="utf-8") as f:
            suggestions = json.load(f)
            return suggestions
    except FileNotFoundError:
        print("No suggestions file found. Using original resume.")
        return []

# Auto-apply enhancement suggestions to resume JSON
def auto_apply_suggestions(resume_json, suggestions):
    """
    Auto-apply all enhancement suggestions to the resume JSON
    """
    updated_resume = resume_json.copy()
    applied_count = 0
    
    print(f"🔄 Applying {len(suggestions)} enhancement suggestions...")
    
    for i, suggestion in enumerate(suggestions, 1):
        suggestion_type = suggestion.get("type")
        target_section = suggestion.get("target_section", "").lower()
        original_snippet = suggestion.get("original_text_snippet", "")
        suggested_text = suggestion.get("suggested_text", "")
        
        print(f"  {i}. [{suggestion_type.upper()}] {target_section}: {suggestion.get('reasoning', '')[:60]}...")
        
        try:
            if suggestion_type == "rephrase":
                if target_section == "summary":
                    updated_resume["summary"] = suggested_text
                    applied_count += 1
                elif target_section == "experience":
                    # Find and replace in experience bullets
                    for exp in updated_resume.get("experience", []):
                        for j, bullet in enumerate(exp.get("bullets", [])):
                            if original_snippet.lower() in bullet.lower():
                                exp["bullets"][j] = suggested_text
                                applied_count += 1
                                break
                
            elif suggestion_type == "add":
                if target_section == "experience":
                    # Add to the first experience entry (usually most recent/relevant)
                    if updated_resume.get("experience"):
                        if "bullets" not in updated_resume["experience"][0]:
                            updated_resume["experience"][0]["bullets"] = []
                        updated_resume["experience"][0]["bullets"].append(suggested_text)
                        applied_count += 1
                elif target_section == "technical skills":
                    # Add new skill category or append to existing
                    if "skills" not in updated_resume:
                        updated_resume["skills"] = {}
                    # Extract the category name from suggested_text (e.g., "SRE & DevOps: ...")
                    if ":" in suggested_text:
                        category, skills_str = suggested_text.split(":", 1)
                        skills_list = [skill.strip() for skill in skills_str.split(",")]
                        updated_resume["skills"][category.strip()] = skills_list
                        applied_count += 1
                
            elif suggestion_type == "quantify":
                # Replace existing text with quantified version
                if target_section == "experience":
                    for exp in updated_resume.get("experience", []):
                        for j, bullet in enumerate(exp.get("bullets", [])):
                            if original_snippet.lower() in bullet.lower():
                                exp["bullets"][j] = suggested_text
                                applied_count += 1
                                break
                                
        except Exception as e:
            print(f"    ❌ Failed to apply suggestion {i}: {e}")
            continue
    
    print(f"✅ Successfully applied {applied_count}/{len(suggestions)} suggestions")
    return updated_resume

# Apply suggestions if available
suggestions = load_suggestions()
if suggestions:
    resume_data = auto_apply_suggestions(resume_data, suggestions)
    print("📝 Resume enhanced with AI suggestions!\n")
else:
    print("📄 Using original resume without enhancements\n")

# Create new Word document
doc = Document()

# ---------- Page Margins ---------- (OPTIMIZED FOR SPACE)
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)    # Reduced from 0.4 to 0.5 for better balance
    section.bottom_margin = Inches(0.5) # Reduced from 0.4 to 0.5 for better balance
    section.left_margin = Inches(0.6)   # Slightly increased for readability
    section.right_margin = Inches(0.6)  # Slightly increased for readability

# --- OPTIMIZED: Maximum space utilization ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)  # Slightly larger than 10 for readability
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.0  # Single line spacing

# ---------- Helper functions ----------

## --- REFINEMENT #2: Add true, clickable hyperlinks ---
def add_hyperlink(paragraph, text, url):
    """
    Adds a hyperlink to a paragraph.
    """
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    run = OxmlElement('w:r')
    run_props = OxmlElement('w:rPr')
    
    # Style the hyperlink to look like a standard link
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000EE') # Blue color
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    
    run_props.append(color)
    run_props.append(underline)
    run.append(run_props)
    
    run.text = text
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    
def add_section(title, extra_space_before=False):
    para = doc.add_paragraph()
    run = para.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(11) # Slightly smaller section title
    
    p_borders = OxmlElement('w:pBdr')
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '4') # Thinner line
    bottom_border.set(qn('w:space'), '1')
    p_borders.append(bottom_border)
    para._p.get_or_add_pPr().append(p_borders)
    
    # --- OPTIMIZED: Minimal spacing for maximum content ---
    if extra_space_before:
        para.paragraph_format.space_before = Pt(8)   # Reduced from 12
    else:
        para.paragraph_format.space_before = Pt(6)   # Reduced from 8
    para.paragraph_format.space_after = Pt(1)       # Reduced from 2

def format_date_range(start, end, education_format=False):
    """Render clean date ranges"""
    start_str = f"{start.get('month', '')} {start.get('year', '')}".strip()
    end_str = f"{end.get('month', '')} {end.get('year', '')}".strip()
    if not start_str and not end_str:
        return ""
    
    # For education: if "Present", show full range; otherwise show only end date
    if education_format:
        if end.get("month", "").lower() == "present":
            return f"{start_str} – Present"
        elif end_str:
            return end_str  # Only show graduation date
        elif start_str:
            return start_str
        else:
            return ""
    
    # For experience and projects: show full date range as before
    if end.get("month", "").lower() == "present":
        return f"{start_str} – Present"
    if not end_str:
        return start_str
    if not start_str:
        return end_str
    if start_str == end_str:
        return start_str
    return f"{start_str} – {end_str}"

            # ---------- Contact Info ----------
if resume_data.get("contactInfo"):
    contact = resume_data["contactInfo"]
    if contact.get("name"):
        name_para = doc.add_paragraph(contact["name"].upper())
        name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = name_para.runs[0]
        run.bold = True
        run.font.size = Pt(16)

    # Create a single paragraph for all contact details
    details_para = doc.add_paragraph()
    details_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add phone number if it exists
    if contact.get("phone"):
        details_para.add_run(contact["phone"])
    
    # Add email with a clickable mailto: link
    if contact.get("email"):
        if details_para.runs: details_para.add_run(" | ")
        add_hyperlink(details_para, contact["email"], f"mailto:{contact['email']}")
    
    # Add other links (LinkedIn, GitHub, etc.)
    for link in contact.get("links", []):
        url = link.get("url", "")
        if url:
            if details_para.runs: details_para.add_run(" | ")
            full_url = url if url.startswith('http') else f"https://{url}"
            add_hyperlink(details_para, url, full_url)

    details_para.paragraph_format.space_after = Pt(8)

# ---------- Summary ----------
if resume_data.get("summary"):
    add_section("Summary")
    p = doc.add_paragraph(resume_data["summary"])
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    
# ---------- Education ----------
if resume_data.get("education"):
    add_section("Education")
    for edu in resume_data["education"]:
        school = edu.get("school", "")
        degree = edu.get("degree", "")
        location = edu.get("location", "")
        start = edu.get("start", {})
        end = edu.get("end", {})
        date_range = format_date_range(start, end, education_format=True)

        # Create paragraph with tab alignment
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.15)
        para.paragraph_format.right_indent = Inches(0.15)
        para.paragraph_format.space_after = Pt(0)

        # Line 1: Degree - School, Location | Dates (date flush right)
        if degree:
            run = para.add_run(degree)
            run.bold = True
            run.font.size = Pt(11)

            if school:
                para.add_run(" - ")
                run = para.add_run(school)
                run.font.size = Pt(10)

                if location:
                    para.add_run(", ")
                    run = para.add_run(location)
                    run.font.size = Pt(10)

        if date_range:
            para.add_run(f"\t{date_range}")
            from docx.enum.text import WD_TAB_ALIGNMENT
            tab_stops = para.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(7.35), WD_TAB_ALIGNMENT.RIGHT)

        # Line 2: Coursework + GPA (GPA always on first line)
        coursework = edu.get("coursework", "")
        gpa = edu.get("gpa", "")

        if coursework or gpa:
            cw_para = doc.add_paragraph()
            cw_para.paragraph_format.left_indent = Inches(0.3)
            cw_para.paragraph_format.right_indent = Inches(0.15)

            if coursework:
                items = [c.strip() for c in coursework.split(",") if c.strip()]
                char_limit = 100  # max characters per line before wrapping
                line = ""
                first_line = True

                for i, item in enumerate(items):
                    part = item + ("," if i < len(items) - 1 else "")
                    if len(line) + len(part) > char_limit:
                        # flush current line
                        run = cw_para.add_run(line.strip())
                        run.font.size = Pt(10)
                        run.italic = True
                        cw_para.add_run().add_break()
                        line = part + " "
                        first_line = False
                    else:
                        line += part + " "

                # flush the last line
                if line.strip():
                    run = cw_para.add_run(line.strip())
                    run.font.size = Pt(10)
                    run.italic = True

            # GPA goes on the **first line only**
            if gpa:
                first_run = cw_para.runs[0] if cw_para.runs else cw_para.add_run("")
                first_run.add_text(f"\t{gpa}")
                from docx.enum.text import WD_TAB_ALIGNMENT
                tab_stops = cw_para.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(7.35), WD_TAB_ALIGNMENT.RIGHT)


# ---------- Experience ----------
if resume_data.get("experience"):
    add_section("Experience")
    for exp in resume_data["experience"]:
        role = exp.get("role", "")
        company = exp.get("company", "")
        location = exp.get("location", "")
        start = exp.get("start", {})
        end = exp.get("end", {})
        date_range = format_date_range(start, end)

        # Create paragraph with tab alignment
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.15)
        para.paragraph_format.right_indent = Inches(0.15)
        para.paragraph_format.space_after = Pt(0)  # Minimal spacing like table version
        
        # Add role and date on same line
        run = para.add_run(role)
        run.bold = True
        run.font.size = Pt(11)  # Make job titles size 11
        
        # Add date aligned to the right using a tab - SAME line as role
        if date_range:
            para.add_run(f"\t{date_range}")
            # Set a right-aligned tab stop at the page's right margin minus our desired spacing
            from docx.enum.text import WD_TAB_ALIGNMENT
            tab_stops = para.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(7.35), WD_TAB_ALIGNMENT.RIGHT)
        
        # Add company on NEW line below role
        if company:
            para.add_run().add_break()
            run = para.add_run(company)
            run.italic = True
            
            # Add location right below the date (aligned to the right)
            if location:
                para.add_run(f"\t{location}")
        elif location:
            # If no company but location exists, add it on new line
            para.add_run().add_break()
            para.add_run(f"\t{location}")

        # Bullets (added after the paragraph)
        for b in exp.get("bullets", []):
            bullet = doc.add_paragraph(b.strip(), style="List Bullet")
            bullet.paragraph_format.left_indent = Inches(0.45)  # Increased from 0.4
            bullet.paragraph_format.right_indent = Inches(0.15)
            bullet.paragraph_format.first_line_indent = Inches(-0.15) # Reduce hanging indent
            bullet.paragraph_format.space_after = Pt(0)
        
        # Add minimal space after the entry - only 2pt (reduced from 4pt)
        if exp != resume_data["experience"][-1]:  # Don't add space after last entry
            doc.paragraphs[-1].paragraph_format.space_after = Pt(2)

# ---------- Projects ----------
if resume_data.get("projects"):
    add_section("Projects")
    for proj in resume_data["projects"]:
        name = proj.get("name", "")
        tech = proj.get("technologies", "")
        start = proj.get("start", {})
        end = proj.get("end", {})
        date_range = format_date_range(start, end)
        url = proj.get("url")  # Check for a project URL

        # Check if this is a bullet-only project (no name, just bullets)
        if not name and proj.get("bullets"):
            # Handle bullet-only projects - just display bullets directly
            for b in proj.get("bullets", []):
                bullet = doc.add_paragraph(b.strip(), style="List Bullet")
                bullet.paragraph_format.left_indent = Inches(0.45)
                bullet.paragraph_format.right_indent = Inches(0.15)
                bullet.paragraph_format.first_line_indent = Inches(-0.15)
                bullet.paragraph_format.space_after = Pt(0)
            
            # Add minimal space after bullet-only entry
            if proj != resume_data["projects"][-1]:
                doc.paragraphs[-1].paragraph_format.space_after = Pt(2)
            continue  # Skip the rest of the processing for this project

        # Regular project processing (with name, tech, dates)
        # Create a simple paragraph with tab alignment
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.15)
        para.paragraph_format.right_indent = Inches(0.15)
        para.paragraph_format.space_after = Pt(0)  # Minimal spacing
        
        # Calculate the content for the first line (project name + technologies)
        project_content = name
        if tech:
            project_content += f" | {tech}"
        
        char_limit = 110  # max characters per line before wrapping
        
        # Check if content fits on one line with the date
        content_with_date_length = len(project_content) + (len(f" {date_range}") if date_range else 0)
        
        if content_with_date_length <= char_limit:
            # Everything fits on one line
            if url:
                add_hyperlink(para, name, url)
                # Make the hyperlinked name bold
                if para.runs:
                    para.runs[-1].bold = True
                    para.runs[-1].font.size = Pt(11)
            else:
                run = para.add_run(name)
                run.bold = True
                run.font.size = Pt(11)

            # Add technologies on the same line
            if tech:
                para.add_run(" | ")
                run = para.add_run(tech)
                run.italic = True
            
            # Add date aligned to the right using a tab
            if date_range:
                para.add_run(f"\t{date_range}")
                from docx.enum.text import WD_TAB_ALIGNMENT
                tab_stops = para.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(7.35), WD_TAB_ALIGNMENT.RIGHT)
        else:
            # Content is too long, split across lines
            # Line 1: Project name + date (date on first line)
            if url:
                add_hyperlink(para, name, url)
                # Make the hyperlinked name bold
                if para.runs:
                    para.runs[-1].bold = True
                    para.runs[-1].font.size = Pt(11)
            else:
                run = para.add_run(name)
                run.bold = True
                run.font.size = Pt(11)
            
            # Add date on the first line
            if date_range:
                para.add_run(f"\t{date_range}")
                from docx.enum.text import WD_TAB_ALIGNMENT
                tab_stops = para.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(7.35), WD_TAB_ALIGNMENT.RIGHT)
            
            # Line 2: Technologies (if they exist)
            if tech:
                para.add_run().add_break()
                run = para.add_run(tech)
                run.italic = True
                run.font.size = Pt(10)
        
        # Bullets
        for b in proj.get("bullets", []):
            bullet = doc.add_paragraph(b.strip(), style="List Bullet")
            bullet.paragraph_format.left_indent = Inches(0.45) # Increased from 0.4
            bullet.paragraph_format.right_indent = Inches(0.15)
            bullet.paragraph_format.first_line_indent = Inches(-0.15) # Reduce hanging indent
            bullet.paragraph_format.space_after = Pt(0)
            
        # Add minimal space after entry - only 2pt (reduced from 4pt)
        if proj != resume_data["projects"][-1]:
            doc.paragraphs[-1].paragraph_format.space_after = Pt(2)
            
# ---------- Publications ----------
if resume_data.get("publications"):
    add_section("Publications")
    for publication in resume_data["publications"]:
        bullet = doc.add_paragraph(publication.strip(), style="List Bullet")
        bullet.paragraph_format.left_indent = Inches(0.45)
        bullet.paragraph_format.right_indent = Inches(0.15)
        bullet.paragraph_format.first_line_indent = Inches(-0.15)
        bullet.paragraph_format.space_after = Pt(0)

# ---------- Achievements ----------
if resume_data.get("achievements"):
    add_section("Achievements")
    for achievement in resume_data["achievements"]:
        bullet = doc.add_paragraph(achievement.strip(), style="List Bullet")
        bullet.paragraph_format.left_indent = Inches(0.45)
        bullet.paragraph_format.right_indent = Inches(0.15)
        bullet.paragraph_format.first_line_indent = Inches(-0.15)
        bullet.paragraph_format.space_after = Pt(0)

# ---------- Technical Skills ----------
if resume_data.get("skills"):
    add_section("Skills")  # Remove extra_space_before to match other sections
    ## --- FIX #5: Correct the gap after each category in skills ---
    # Build the entire skills block as a single string, then add it
    skills_text = []
    if isinstance(resume_data["skills"], dict):
        for category, items in resume_data["skills"].items():
            # Make category names bold
            skills_text.append(f"**{category}**: {', '.join(items)}")
    
    # Add the combined text as one paragraph with line breaks
    skills_para = doc.add_paragraph()
    skills_para.paragraph_format.left_indent = Inches(0.15)
    skills_para.paragraph_format.right_indent = Inches(0.15)
    for i, skill_line in enumerate(skills_text):
        if i > 0:
            skills_para.add_run().add_break()
        
        # Split by ** to handle bold formatting
        parts = skill_line.split('**')
        for j, part in enumerate(parts):
            if j % 2 == 1:  # Odd indices are between ** markers (bold)
                run = skills_para.add_run(part)
                run.bold = True
            else:
                skills_para.add_run(part)
    
# Save DOCX
output_file = "Enhanced_Resume.docx"
doc.save(output_file)
print(f"✅ Resume saved to {output_file}")